"""
loader.py — Load legal documents from a folder.
Supports PDF (via PyMuPDF), TXT, DOCX.
Each Document carries metadata: source, category, filename, page.
"""

from __future__ import annotations

import os
from pathlib import Path
from typing import List

from langchain_core.documents import Document

SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".docx", ".md"}


def _load_pdf(path: Path) -> List[Document]:
    import fitz  # PyMuPDF
    docs: List[Document] = []
    pdf = fitz.open(str(path))
    for page_num, page in enumerate(pdf):
        text = page.get_text()
        if text.strip():
            docs.append(Document(
                page_content=text,
                metadata={
                    "source": str(path),
                    "filename": path.name,
                    "page": page_num + 1,
                    "category": _infer_category(path),
                },
            ))
    pdf.close()
    return docs


def _load_txt(path: Path) -> List[Document]:
    text = path.read_text(encoding="utf-8", errors="ignore")
    return [Document(
        page_content=text,
        metadata={
            "source": str(path),
            "filename": path.name,
            "page": 1,
            "category": _infer_category(path),
        },
    )]


def _load_docx(path: Path) -> List[Document]:
    import docx
    doc = docx.Document(str(path))
    text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    return [Document(
        page_content=text,
        metadata={
            "source": str(path),
            "filename": path.name,
            "page": 1,
            "category": _infer_category(path),
        },
    )]


def _infer_category(path: Path) -> str:
    """Infer category from folder structure: case_law / contracts / general."""
    parts = [p.lower() for p in path.parts]
    if "case_law" in parts:
        return "case_law"
    if "contracts" in parts:
        return "contracts"
    return "general"


_LOADERS = {
    ".pdf": _load_pdf,
    ".txt": _load_txt,
    ".docx": _load_docx,
    ".md": _load_txt,
}


def load_documents_from_folder(folder: str) -> List[Document]:
    """Recursively load all supported legal documents from *folder*."""
    docs: List[Document] = []
    for root, _, files in os.walk(folder):
        for fname in sorted(files):
            path = Path(root) / fname
            ext = path.suffix.lower()
            if ext in _LOADERS:
                try:
                    docs.extend(_LOADERS[ext](path))
                except Exception as exc:
                    print(f"[loader] Skipping {path}: {exc}")
    return docs
